"""
DOCX 解析器 — 基于 python-docx。

从 Word 文档中提取段落文本、表格、图片，返回 Document 列表。
"""

import hashlib
import os
from typing import List, Optional

from langchain_core.documents import Document
from tqdm import tqdm

from kb.offline.scripts.file_processors.base_parser import BaseParser


class DOCXParser(BaseParser):
    """Word (.docx) 格式解析器。"""

    def supported_extensions(self) -> List[str]:
        return [".docx"]

    def parse(
        self,
        file_path: str,
        extract_images: bool = True,
        include_tables: bool = True,
        **kwargs,
    ) -> List[Document]:
        """
        解析 DOCX 文件。

        Args:
            file_path: .docx 文件路径
            extract_images: 是否提取并保存图片
            include_tables: 是否将表格内容作为独立文档块

        Returns:
            List[Document]: 解析结果
        """
        try:
            from docx import Document as DocxDocument
            from docx.oxml.ns import qn
        except ImportError:
            raise ImportError(
                "请安装 python-docx: pip install python-docx"
            )

        doc = DocxDocument(file_path)
        raw_docs: List[Document] = []
        images_info = []

        # 提取图片（DOCX 中的图片内嵌在文档中）
        if extract_images:
            images_info = self._extract_images(doc, file_path)

        # 按段落提取文本，组合成合理大小的文档块
        current_text = []
        current_paragraphs = 0
        max_paragraphs_per_chunk = 20  # 每块最多 20 段落

        for para in tqdm(doc.paragraphs, desc="解析 DOCX 段落"):
            text = para.text.strip()
            if not text:
                continue

            current_text.append(text)
            current_paragraphs += 1

            # 遇到标题或达到段落数上限时切分
            if (
                para.style and para.style.name and "heading" in para.style.name.lower()
            ) or current_paragraphs >= max_paragraphs_per_chunk:
                self._flush_paragraphs(
                    current_text, current_paragraphs, file_path, raw_docs, images_info
                )
                current_text = []
                current_paragraphs = 0

        # 处理剩余段落
        if current_text:
            self._flush_paragraphs(
                current_text, current_paragraphs, file_path, raw_docs, images_info
            )

        # 提取表格内容
        if include_tables and doc.tables:
            self._parse_tables(doc, file_path, raw_docs)

        print(f"[DOCXParser] 解析完成: {file_path} → {len(raw_docs)} 个文档块")
        return raw_docs

    def _flush_paragraphs(
        self,
        text_parts: List[str],
        para_count: int,
        file_path: str,
        raw_docs: List[Document],
        images_info: List[dict],
    ) -> None:
        """将缓存的段落写入一个 Document 块。"""
        combined = "\n".join(text_parts)
        unique_id = hashlib.md5(combined.encode("utf-8")).hexdigest()
        metadata = {
            "unique_id": unique_id,
            "source": file_path,
            "format": "docx",
            "images_info": images_info,
        }
        raw_docs.append(Document(page_content=combined, metadata=metadata))

    def _parse_tables(
        self, doc, file_path: str, raw_docs: List[Document]
    ) -> None:
        """提取 DOCX 中的表格内容。"""
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_text = "\n".join(rows)
            if table_text.strip():
                unique_id = hashlib.md5(table_text.encode("utf-8")).hexdigest()
                metadata = {
                    "unique_id": unique_id,
                    "source": file_path,
                    "format": "docx",
                    "table_index": table_idx,
                }
                raw_docs.append(Document(page_content=table_text, metadata=metadata))

    def _extract_images(self, doc, file_path: str) -> List[dict]:
        """
        提取 DOCX 中的内嵌图片并保存。

        注意: python-docx 对图片的提取能力有限，
        复杂场景建议使用 python-pptx 方式处理或保持原图。
        """
        images_info = []
        try:
            # 获取文档保存目录
            doc_dir = os.path.dirname(file_path)
            image_subdir = os.path.join(doc_dir, ".images")
            os.makedirs(image_subdir, exist_ok=True)

            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    img_data = rel.target_part.blob
                    img_ext = os.path.splitext(rel.target_ref)[1] or ".png"
                    img_name = f"docx_img_{len(images_info) + 1}{img_ext}"
                    img_path = os.path.join(image_subdir, img_name)
                    with open(img_path, "wb") as f:
                        f.write(img_data)
                    images_info.append({
                        "image_path": img_path,
                        "page": 0,
                        "title": "",
                    })
        except Exception as e:
            print(f"[DOCXParser] 图片提取警告: {e}")

        return images_info


# ---- 便利函数 ----
def load_docx(file_path: str, **kwargs) -> List[Document]:
    """便利函数，可直接调用。"""
    return DOCXParser().parse(file_path, **kwargs)
